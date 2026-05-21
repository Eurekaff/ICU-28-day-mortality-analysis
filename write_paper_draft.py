from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .settings import DOCS_DIR, FIGURES_DIR, PROCESSED_DIR, TABLES_DIR


OUTPUT_NAME = "基于MIMIC-IV的阿尔茨海默病ICU患者延长ICU住院风险预测研究_论文初稿.docx"
OUTPUT_PATH = DOCS_DIR / OUTPUT_NAME


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_font(paragraph, size: int = 12, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        run.bold = bold


def add_body_paragraph(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)
    return para


def add_center_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    return para


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True
        run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], headers: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        set_cell_shading(header_cells[idx], "D9EAF7")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.3f}"
            set_cell_text(cells[idx], str(value))

    doc.add_paragraph()


def fmt_metric(value: float) -> str:
    return f"{value:.3f}"


def load_results() -> dict[str, pd.DataFrame]:
    return {
        "summary": pd.read_csv(TABLES_DIR / "prolonged_icu_los_dataset_summary.csv"),
        "performance": pd.read_csv(TABLES_DIR / "prolonged_icu_los_table2_model_performance.csv"),
        "ci": pd.read_csv(TABLES_DIR / "prolonged_icu_los_table4_best_model_bootstrap_ci.csv"),
        "importance": pd.read_csv(TABLES_DIR / "prolonged_icu_los_feature_importance.csv"),
        "confusion": pd.read_csv(TABLES_DIR / "prolonged_icu_los_confusion_matrix.csv"),
        "meta": pd.read_csv(PROCESSED_DIR / "prolonged_icu_los_best_model_meta.csv"),
    }


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def build_docx() -> Path:
    results = load_results()
    summary = results["summary"].iloc[0]
    perf = results["performance"].copy()
    ci = results["ci"].copy()
    importance = results["importance"].head(10).copy()
    confusion = results["confusion"].copy()
    best = results["meta"].iloc[0]

    doc = Document()
    configure_document(doc)

    add_center_paragraph(doc, "本  科  毕  业  论  文 （设 计）", size=18, bold=True)
    add_center_paragraph(doc, "（主修专业）", size=14)
    doc.add_paragraph()
    add_center_paragraph(doc, "基于 MIMIC-IV 的阿尔茨海默病 ICU 患者延长 ICU 住院风险预测研究", size=16, bold=True)
    add_center_paragraph(
        doc,
        "Prediction of Prolonged ICU Length of Stay in ICU Patients with Alzheimer's Disease Based on MIMIC-IV",
        size=12,
    )
    doc.add_paragraph()
    for line in [
        "姓    名：方明正",
        "学    号：37220222203586",
        "学    院：信息学院",
        "专    业：软件工程",
        "年    级：2022级",
        "校内指导教师：黄晨曦   助理教授",
    ]:
        add_center_paragraph(doc, line, size=12)
    doc.add_paragraph()
    add_center_paragraph(doc, "二〇二六年 五月", size=12)
    doc.add_page_break()

    add_center_paragraph(doc, "摘　要", size=16, bold=True)
    for text in [
        "阿尔茨海默病（Alzheimer's disease，AD）患者进入重症监护病房（ICU）后，常伴随年龄较高、基础疾病复杂、功能状态较差和照护需求较高等特点。除短期死亡外，ICU 住院时间延长也是反映疾病负担、医疗资源占用和照护复杂度的重要结局。对于 AD-ICU 患者，若能在入 ICU 早期识别延长 ICU 住院风险，有助于辅助临床资源配置和患者管理。",
        "本文基于 MIMIC-IV v3.1 数据库，在既有 AD-ICU 队列和首个 24 小时特征工程基础上，构建延长 ICU 住院风险预测模型。研究对象为成年首次 ICU 住院且被识别为阿尔茨海默病的患者，并排除 ICU 停留不足 24 小时或 ICU 住院时长缺失者。延长 ICU 住院默认定义为 ICU 住院天数大于 7 天，标签阈值在程序中设置为可配置项，以便后续采用中位数、上四分位数或其他临床定义。",
        "本研究纳入 579 例 AD-ICU 患者，其中 62 例发生延长 ICU 住院，事件率为 10.71%。基于人口学信息、首个 ICU 收治单元、生命体征和实验室指标等 43 个候选特征，采用分层训练测试划分和 5 折分层交叉验证，比较 Logistic 回归、决策树、随机森林、极端随机树、直方图梯度提升、XGBoost、LightGBM 和 CatBoost 等模型。模型评价指标包括 AUROC、AUPRC、准确率、精确率、召回率、F1 值和 Brier score，并补充混淆矩阵、校准曲线和特征重要性分析。",
        f"结果显示，CatBoost 在测试集上取得最高 AUROC（{fmt_metric(best['best_model_test_auroc'])}），AUPRC 为 {fmt_metric(perf.loc[perf['model'] == 'CatBoost', 'test_auprc'].iloc[0])}，Brier score 为 {fmt_metric(perf.loc[perf['model'] == 'CatBoost', 'test_brier'].iloc[0])}。在以训练集 F1 最优阈值进行分类时，部分集成模型倾向于给出较保守的阳性预测，提示该任务受到明显类别不平衡和样本量限制影响。特征重要性结果显示，首个 24 小时血氧饱和度均值、ICU 收治单元、年龄、种族、血压和呼吸频率等变量对模型判断贡献较高。",
        "综上，基于首个 24 小时结构化临床数据构建 AD-ICU 患者延长 ICU 住院风险预测模型具有一定可行性，但当前样本规模和事件数量有限，模型区分能力仍处于探索阶段。研究结果可为后续围绕阿尔茨海默病 ICU 患者资源消耗、住院管理和风险分层的研究提供方法基础。",
    ]:
        add_body_paragraph(doc, text)
    add_body_paragraph(doc, "关键词：阿尔茨海默病；重症监护病房；住院时长；风险预测；机器学习；MIMIC-IV")
    doc.add_page_break()

    add_center_paragraph(doc, "Abstract", size=16, bold=True)
    for text in [
        "Patients with Alzheimer's disease admitted to the intensive care unit usually present with older age, complex comorbidities, impaired functional status, and high care needs. In addition to short-term mortality, prolonged ICU length of stay is an important outcome reflecting disease burden, resource consumption, and care complexity.",
        "Using the MIMIC-IV v3.1 database, this study developed machine learning models for predicting prolonged ICU length of stay among AD-ICU patients. Adult patients with their first ICU stay and identified Alzheimer's disease were included. Prolonged ICU stay was defined as ICU length of stay greater than 7 days by default, and the threshold was implemented as a configurable parameter.",
        f"A total of {int(summary['n_total'])} AD-ICU patients were included, among whom {int(summary['n_events'])} experienced prolonged ICU stay. Logistic regression, decision tree, random forest, extra trees, histogram gradient boosting, XGBoost, LightGBM, and CatBoost were compared using stratified train-test split and stratified cross-validation. CatBoost achieved the highest test AUROC of {fmt_metric(best['best_model_test_auroc'])}. However, classification performance at the selected threshold was limited, suggesting the influence of class imbalance and small event counts.",
        "The findings indicate that early structured clinical data can provide useful information for risk stratification of prolonged ICU stay in AD-ICU patients, but further validation with larger and external cohorts is needed.",
    ]:
        add_body_paragraph(doc, text)
    add_body_paragraph(doc, "Key words: Alzheimer's disease; intensive care unit; length of stay; risk prediction; machine learning; MIMIC-IV")
    doc.add_page_break()

    add_center_paragraph(doc, "目 录", size=16, bold=True)
    toc_lines = [
        "第一章 绪论",
        "1.1 研究背景",
        "1.2 研究现状与问题提出",
        "1.3 研究目的与主要研究问题",
        "1.4 研究内容与技术路线",
        "第二章 资料来源与研究方法",
        "2.1 数据来源",
        "2.2 研究对象筛选",
        "2.3 结局定义与候选特征",
        "2.4 数据预处理与模型构建",
        "2.5 模型评估方法",
        "第三章 实验结果",
        "3.1 队列构建与标签分布",
        "3.2 模型性能比较",
        "3.3 阈值分类结果与校准分析",
        "3.4 特征重要性分析",
        "第四章 讨论",
        "第五章 总结与展望",
        "参考文献",
    ]
    for line in toc_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.8 if line[0].isdigit() else 0)
        set_paragraph_font(p, 12)
    doc.add_page_break()

    add_heading(doc, "第一章 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    for text in [
        "阿尔茨海默病是老年人群中最常见的神经退行性疾病之一，随着人口老龄化持续加深，合并 AD 的住院患者数量逐渐增加。AD 患者一旦进入 ICU，往往不仅面临原发急危重症的影响，还受到认知功能减退、基础功能状态下降、沟通困难和照护依赖等因素影响，因此其 ICU 管理过程通常比一般患者更复杂。",
        "在重症医学研究中，死亡风险是最常见的结局指标，但住院时长同样具有重要意义。ICU 住院时间延长通常意味着病情恢复缓慢、并发症风险增加、医疗资源占用增加和后续照护压力加重。对于 AD 患者而言，延长 ICU 住院还可能进一步影响谵妄、功能退化、出院去向和家庭照护负担。因此，在 AD-ICU 患者中开展延长 ICU 住院风险预测具有现实意义。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.2 研究现状与问题提出", 2)
    for text in [
        "已有研究从不同角度讨论了痴呆或 AD 患者在 ICU 场景中的预后表现，提示该人群可能具有更高死亡风险、更长住院时间和更高资源消耗。与此同时，机器学习方法已被广泛应用于 ICU 死亡、再入院和住院时长预测任务，并显示出处理多维结构化电子病历数据的优势。",
        "不过，现有研究仍存在两方面不足。第一，许多研究以全体 ICU 患者为对象，对 AD 患者这类特殊亚群的延长住院风险关注不足。第二，部分研究虽然涉及 ICU 住院时长预测，但往往未将标签定义、队列筛选、特征工程和模型评价组织成面向 AD-ICU 患者的独立研究流程。基于此，本文在既有 AD-ICU 队列和特征工程基础上，进一步构建延长 ICU 住院风险预测模型。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.3 研究目的与主要研究问题", 2)
    for text in [
        "本文的研究目的，是基于 MIMIC-IV 数据库构建 AD-ICU 患者延长 ICU 住院风险预测模型，并分析首个 24 小时结构化临床数据对该任务的预测价值。",
        "本文主要回答三个问题：第一，在 AD-ICU 患者中，采用 ICU 住院天数大于 7 天作为延长 ICU 住院定义时，事件分布如何；第二，基于入 ICU 后首个 24 小时临床信息，哪些机器学习模型能够更好地区分延长住院风险；第三，模型输出在阈值分类、校准和特征重要性层面表现如何。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.4 研究内容与技术路线", 2)
    for text in [
        "本文首先复用既有 AD-ICU 队列构建和首个 24 小时特征提取流程，得到包含人口学、收治单元、生命体征和实验室指标的数据集。随后以 ICU 住院天数大于 7 天构造二分类标签，并将阈值保留为配置项。最后，比较多类机器学习模型的预测表现，输出模型性能表、混淆矩阵、校准曲线、ROC 曲线、PR 曲线和特征重要性结果。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "第二章 资料来源与研究方法", 1)
    add_heading(doc, "2.1 数据来源", 2)
    for text in [
        "本研究数据来源于 MIMIC-IV v3.1 数据库。MIMIC-IV 是公开、去标识化的重症医学电子病历数据库，包含住院、ICU、实验室检查、生命体征、诊断编码和结局等信息。本文主要使用住院模块和 ICU 模块中的结构化数据，并基于既有项目流程完成数据读取、队列构建和特征整理。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.2 研究对象筛选", 2)
    for text in [
        "基础队列为成年 ICU 患者的首次 ICU 住院记录。阿尔茨海默病识别主要基于诊断编码标记。预测模型队列进一步限定为 AD 标记为阳性、ICU 住院时长不缺失且 ICU 停留时间不少于 24 小时的患者。该筛选规则与既有 AD-ICU 建模流程保持一致，以保证首个 24 小时特征具有明确含义。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.3 结局定义与候选特征", 2)
    for text in [
        "主要结局为延长 ICU 住院。本文默认将 ICU 住院天数大于 7 天定义为延长 ICU 住院，即 prolonged_icu_los = 1 if icu_los_days > 7 else 0。该阈值在代码中设置为可配置项，后续可根据研究需要改为中位数、上四分位数或其他固定天数。",
        "候选特征包括年龄、性别、种族、首个 ICU 收治单元，以及首个 24 小时内生命体征和实验室指标的首值、均值、最小值和最大值等摘要特征。实际进入模型的数据集中共有 40 个数值特征和 3 个分类特征。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.4 数据预处理与模型构建", 2)
    for text in [
        "数值变量采用中位数填补缺失值，并进行标准化处理；分类变量采用众数填补，并通过 one-hot 编码转换为模型可用的特征。训练集和测试集采用分层随机划分，测试集比例为 20%。训练过程中使用 5 折分层交叉验证进行超参数选择，评价标准为 AUROC。",
        "本文比较 Logistic 回归、决策树、随机森林、极端随机树、直方图梯度提升、XGBoost、LightGBM 和 CatBoost 共 8 类模型。对于类别不平衡问题，优先使用模型自身的 class_weight 或 auto_class_weights 等机制进行处理，不额外引入复杂采样依赖。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.5 模型评估方法", 2)
    for text in [
        "模型评价包括概率排序能力、分类效果和概率校准三个层面。概率排序能力主要采用 AUROC 和 AUPRC；分类效果采用准确率、精确率、召回率和 F1 值；概率输出质量采用 Brier score。分类阈值根据训练集上 F1 值最优原则确定，并在测试集上报告混淆矩阵。对于最优模型，进一步采用 bootstrap 方法估计主要指标的 95% 置信区间。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "第三章 实验结果", 1)
    add_heading(doc, "3.1 队列构建与标签分布", 2)
    add_body_paragraph(
        doc,
        f"按照研究对象筛选规则，本研究最终纳入 {int(summary['n_total'])} 例 AD-ICU 患者。其中 ICU 住院天数大于 7 天者 {int(summary['n_events'])} 例，事件率为 {summary['event_rate'] * 100:.2f}%。模型训练使用 {int(best['n_train'])} 例，测试集使用 {int(best['n_test'])} 例，其中测试集延长住院事件数为 {int(best['n_test_events'])} 例。",
    )
    add_caption(doc, "表 3-1 延长 ICU 住院预测数据集概况")
    add_df_table(
        doc,
        pd.DataFrame([summary]),
        ["threshold_days", "threshold_hours", "n_total", "n_events", "event_rate", "n_numeric_features", "n_categorical_features"],
        ["阈值(天)", "阈值(小时)", "样本量", "事件数", "事件率", "数值特征", "分类特征"],
    )

    add_heading(doc, "3.2 模型性能比较", 2)
    add_body_paragraph(
        doc,
        f"8 类模型均完成训练和测试。按测试集 AUROC 排序，CatBoost 表现最好，AUROC 为 {fmt_metric(best['best_model_test_auroc'])}；DecisionTree 排名第二，AUROC 为 {fmt_metric(perf.loc[perf['model'] == 'DecisionTree', 'test_auroc'].iloc[0])}。从 AUPRC 看，DecisionTree 的 AUPRC 略高于 CatBoost，但整体数值仍较低，说明在事件率较低的情况下，阳性病例识别仍具有较大难度。",
    )
    add_caption(doc, "表 3-2 不同模型在测试集上的预测性能")
    add_df_table(
        doc,
        perf,
        ["model", "test_auroc", "test_auprc", "test_accuracy", "test_precision", "test_recall", "test_f1", "test_brier"],
        ["模型", "AUROC", "AUPRC", "准确率", "精确率", "召回率", "F1", "Brier"],
    )

    for image_name, caption in [
        ("prolonged_icu_los_roc.png", "图 3-1 不同模型 ROC 曲线比较"),
        ("prolonged_icu_los_pr.png", "图 3-2 不同模型 PR 曲线比较"),
    ]:
        image_path = FIGURES_DIR / image_name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(5.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, caption)

    add_heading(doc, "3.3 阈值分类结果与校准分析", 2)
    add_body_paragraph(
        doc,
        "在分类阈值层面，部分集成模型虽然具有相对较高的 AUROC 或较低的 Brier score，但在训练集 F1 最优阈值迁移到测试集后倾向于给出保守预测，测试集中未识别出阳性病例。DecisionTree 和 Logistic 回归能够识别部分阳性病例，但伴随较多假阳性。这说明在当前事件数较少的条件下，仅依赖单一阈值进行阳性判别并不稳定，更适合将模型输出作为风险排序参考。",
    )
    add_caption(doc, "表 3-3 不同模型在测试集上的混淆矩阵")
    add_df_table(
        doc,
        confusion,
        ["model", "threshold", "tn", "fp", "fn", "tp"],
        ["模型", "阈值", "TN", "FP", "FN", "TP"],
    )
    image_path = FIGURES_DIR / "prolonged_icu_los_calibration.png"
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(5.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "图 3-3 最优模型校准曲线")

    add_caption(doc, "表 3-4 最优模型主要指标 bootstrap 95% 置信区间")
    add_df_table(
        doc,
        ci,
        ["metric", "mean", "lcl_95", "ucl_95"],
        ["指标", "均值", "95%CI下限", "95%CI上限"],
    )

    add_heading(doc, "3.4 特征重要性分析", 2)
    add_body_paragraph(
        doc,
        "对最优 CatBoost 模型进行特征重要性分析后发现，排名靠前的变量包括首个 24 小时血氧饱和度均值、首个 ICU 收治单元、年龄、种族、舒张压均值、收缩压最小值、平均动脉压均值和呼吸频率均值等。这些变量既反映患者基础状态，也反映入 ICU 早期循环和呼吸状态，与延长 ICU 住院风险具有一定临床解释一致性。",
    )
    add_caption(doc, "表 3-5 最优模型前 10 位重要特征")
    add_df_table(
        doc,
        importance,
        ["feature", "importance"],
        ["特征", "重要性"],
    )

    add_heading(doc, "第四章 讨论", 1)
    for text in [
        "本文结果显示，AD-ICU 患者中延长 ICU 住院事件并不少见，但以大于 7 天作为阈值时事件率约为 10.71%，属于相对不平衡的二分类任务。在这种背景下，Accuracy 容易受到阴性样本占比影响，不能单独作为模型优劣判断依据。因此，本文同时报告 AUROC、AUPRC、Brier score 和混淆矩阵，以更全面地描述模型表现。",
        "从模型比较看，CatBoost 的 AUROC 最高，说明其对患者风险排序具有一定优势；但在固定阈值下，CatBoost 未在测试集中预测出阳性病例。DecisionTree 虽然 AUROC 略低，但在测试集中识别出 4 例真实延长住院患者，召回率为 33.33%，F1 值为 0.258。这提示，在小样本不平衡任务中，“最佳排序模型”和“最佳阈值分类模型”可能并不完全一致，实际应用中需要根据使用场景选择模型和阈值。",
        "从特征重要性看，呼吸状态、循环状态、年龄和 ICU 收治背景是模型判断的重要信息来源。血氧饱和度、呼吸频率、血压和平均动脉压等变量反映患者入 ICU 早期生理稳定性；ICU 收治单元和种族等变量可能同时包含疾病类型、收治路径和群体差异信息。上述结果提示，延长 ICU 住院风险并非只由单一严重程度指标决定，而是多维临床信息共同作用的结果。",
        "本研究仍存在局限。首先，研究基于单中心公开数据库，外部推广性仍需验证。其次，AD 识别依赖诊断编码，无法进一步区分 AD 严重程度。再次，延长 ICU 住院采用大于 7 天作为默认阈值，虽然具有直观解释性，但不同研究场景可能采用中位数、上四分位数或其他定义。最后，当前事件数较少，模型阈值分类结果不稳定，后续可结合更多样本、外部验证和代价敏感阈值策略进一步优化。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "第五章 总结与展望", 1)
    for text in [
        "本文基于 MIMIC-IV v3.1 数据库，在既有 AD-ICU 队列和首个 24 小时特征工程基础上，构建了阿尔茨海默病 ICU 患者延长 ICU 住院风险预测模型。研究将 ICU 住院天数大于 7 天定义为延长 ICU 住院，并将该阈值设置为可配置项。最终纳入 579 例患者，其中 62 例发生延长 ICU 住院。",
        "实验比较了 8 类机器学习模型。CatBoost 在测试集上取得最高 AUROC，提示早期结构化数据对延长住院风险排序具有一定价值；但受类别不平衡和事件数限制，阈值分类表现仍不稳定。总体而言，本研究为 AD-ICU 患者延长 ICU 住院风险预测提供了可运行、可复现的实验流程和初步结果。",
        "后续研究可从三个方向继续推进：第一，尝试不同延长住院定义，并比较固定阈值、中位数和上四分位数标签下的模型稳定性；第二，引入治疗干预、并发症、用药和动态轨迹特征，进一步提高模型对住院过程的刻画能力；第三，在外部数据库或不同医院队列中开展验证，以评估模型的泛化能力和实际应用价值。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] Johnson A, Bulgarelli L, Pollard T, et al. MIMIC-IV [DS]. PhysioNet, 2024.",
        "[2] Johnson A E W, Bulgarelli L, Shen L, et al. MIMIC-IV, a freely accessible electronic health record dataset [J]. Scientific Data, 2023, 10(1): 1.",
        "[3] Yorganci E, Sleeman K E, Sampson E L, et al. Survival and critical care use among people with dementia in a large English cohort [J]. Age and Ageing, 2023, 52(9): afad157.",
        "[4] Dziegielewski C, Fernando S M, Milani C, et al. Outcomes and cost analysis of patients with dementia in the intensive care unit: a population-based cohort study [J]. BMC Health Services Research, 2023, 23(1): 1124.",
        "[5] Davis-Ajami M L, Chang C-H, Gupta S, et al. Mortality and discharge location of intensive care patients with Alzheimer disease and related dementia [J]. American Journal of Critical Care, 2023, 32(4): 249-255.",
        "[6] Zhu B, Chen X, Li W, et al. Effect of Alzheimer disease on prognosis of intensive care unit patients: a propensity score matching analysis [J]. Medical Science Monitor, 2022, 28: e936550.",
        "[7] Liu H, Liang Q, Yang Y, et al. Impact of mechanical ventilation on clinical outcomes in ICU-admitted Alzheimer's disease patients: a retrospective cohort study [J]. Frontiers in Public Health, 2024, 12: 1368508.",
        "[8] Olang O, Mohseni S, Shahabinezhad A, et al. Artificial intelligence-based models for prediction of mortality in ICU patients: a scoping review [J]. Journal of Intensive Care Medicine, 2025, 40(12): 1240-1246.",
        "[9] Deng Y, Liu S, Wang Z, et al. Explainable time-series deep learning models for the prediction of mortality, prolonged length of stay and 30-day readmission in intensive care patients [J]. Frontiers in Medicine, 2022, 9: 933037.",
        "[10] Lundberg S M, Lee S-I. A unified approach to interpreting model predictions [J]. Advances in Neural Information Processing Systems, 2017, 30.",
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.line_spacing = 1.2
        set_paragraph_font(p, 10)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build_docx()
    print(out)
